from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"D:\工作文档\00_进行中\艾维美术馆")
SOURCE = ROOT / "00_工作台" / "近期工作计划" / "202607-会员制艺术阅读会方案.md"
OUTPUT = ROOT / "00_工作台" / "近期工作计划" / "202607-会员制艺术阅读会方案.docx"
LOGO = ROOT / "00_工作台" / "近期工作计划" / "assets" / "aiwei-brand" / "logo.png"

PINE = "143E3B"
RED = "E30016"
INK = "1F2927"
MUTED = "66736F"
MIST = "EAF0ED"
LINE = "C9D7D1"
PAPER = "F7F9F8"
FONT = "Microsoft YaHei"
CONTENT_WIDTH = 9360


def set_run_font(run, size=None, color=None, bold=None, italic=None):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for grid_col, width in zip(grid.gridCol_lst, widths):
        grid_col.set(qn("w:w"), str(width))
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    rel_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), RED)
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    east_asia = OxmlElement("w:rFonts")
    east_asia.set(qn("w:eastAsia"), FONT)
    r_pr.append(east_asia)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def apply_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.33

    for name, size, color, before, after in [
        ("Heading 1", 16, PINE, 18, 10),
        ("Heading 2", 13, PINE, 12, 6),
        ("Heading 3", 11.5, RED, 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Source" not in [s.name for s in doc.styles]:
        source = doc.styles.add_style("Source", WD_STYLE_TYPE.PARAGRAPH)
        source.base_style = normal
        source.font.name = FONT
        source._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        source.font.size = Pt(9)
        source.font.color.rgb = RGBColor.from_string(MUTED)
        source.paragraph_format.space_after = Pt(3)


def add_cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(38)
    run = p.add_run()
    run.add_picture(str(LOGO), width=Inches(2.25))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(64)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run("项目方案")
    set_run_font(run, 11, RED, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("会员制艺术阅读会")
    set_run_font(run, 28, PINE, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(42)
    run = p.add_run("在美术馆发生的持续艺术阅读社群")
    set_run_font(run, 13, MUTED)

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    set_table_geometry(table, [2600, 6760])
    rows = [("项目状态", "方案已形成，待立项确认"), ("首期形式", "创始试读季：连续 12 个周六"), ("更新时间", "2026-07-23")]
    for row, (label, value) in zip(table.rows, rows):
        set_cell_shading(row.cells[0], PINE)
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        r = row.cells[0].paragraphs[0].add_run(label)
        set_run_font(r, 10, "FFFFFF", bold=True)
        r = row.cells[1].paragraphs[0].add_run(value)
        set_run_font(r, 10, INK)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def add_table(doc, rows):
    cols = len(rows[0])
    weights = [max(3, max(len(r[i]) for r in rows)) for i in range(cols)]
    total = sum(weights)
    widths = [round(CONTENT_WIDTH * w / total) for w in weights]
    widths[-1] += CONTENT_WIDTH - sum(widths)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    set_table_geometry(table, widths)
    for r_i, row_data in enumerate(rows):
        for c_i, value in enumerate(row_data):
            cell = table.cell(r_i, c_i)
            if r_i == 0:
                set_cell_shading(cell, PINE)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(value)
            set_run_font(run, 9.2, "FFFFFF" if r_i == 0 else INK, bold=(r_i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def parse_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        cells = [c.strip() for c in lines[i].strip().strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            rows.append(cells)
        i += 1
    return rows, i


def add_markdown_paragraph(doc, text, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(8 if style is None else 4)
    # Render markdown links as clickable Word hyperlinks.
    parts = re.split(r"(\[[^\]]+\]\([^\)]+\))", text)
    for part in parts:
        match = re.fullmatch(r"\[([^\]]+)\]\(([^\)]+)\)", part)
        if match:
            add_hyperlink(p, match.group(1), match.group(2))
        else:
            plain = re.sub(r"`([^`]+)`", r"\1", part)
            run = p.add_run(plain)
            set_run_font(run, 10.5, INK)
    return p


def add_body(doc, markdown):
    lines = markdown.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            i += 1
            continue
        if line.startswith("# "):
            i += 1
            continue
        if line.startswith("## "):
            doc.add_paragraph(line[3:], style="Heading 1")
        elif line.startswith("### "):
            doc.add_paragraph(line[4:], style="Heading 2")
        elif line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table(doc, rows)
            continue
        elif re.match(r"^- ", line):
            add_markdown_paragraph(doc, line[2:], style="List Bullet")
        elif re.match(r"^\d+\. ", line):
            add_markdown_paragraph(doc, re.sub(r"^\d+\. ", "", line), style="List Number")
        else:
            p = add_markdown_paragraph(doc, line)
            if line.startswith("更新时间：") or line.startswith("项目状态：") or line.startswith("项目定位："):
                for run in p.runs:
                    set_run_font(run, 10, MUTED)
        i += 1


def set_header_footer(section):
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("艾维美术馆  |  会员制艺术阅读会方案")
    set_run_font(run, 8.5, MUTED)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("艾维美术馆  ·  2026  ·  第 ")
    set_run_font(run, 8.5, MUTED)
    add_page_number(p)
    run = p.add_run(" 页")
    set_run_font(run, 8.5, MUTED)


def main():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    apply_styles(doc)
    set_header_footer(section)
    add_cover(doc)
    add_body(doc, SOURCE.read_text(encoding="utf-8"))
    doc.core_properties.title = "艾维美术馆会员制艺术阅读会方案"
    doc.core_properties.subject = "会员制艺术阅读会项目方案"
    doc.core_properties.author = "艾维美术馆"
    doc.core_properties.comments = "由项目方案 Markdown 转换。"
    doc.save(OUTPUT)
    verified = Document(OUTPUT)
    headings = [p.text for p in verified.paragraphs if p.style.name.startswith("Heading")]
    print({
        "output": str(OUTPUT),
        "paragraphs": len(verified.paragraphs),
        "tables": len(verified.tables),
        "headings": len(headings),
        "has_logo": any("image" in rel.target_ref for rel in verified.part.rels.values()),
    })


if __name__ == "__main__":
    main()
